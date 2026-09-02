#!/usr/bin/env python3
"""Audit structural invariants in the reviewed SCI manuscript template."""

from __future__ import annotations

import argparse
import json
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_TEMPLATE = ROOT / "assets" / "SCI_manuscript_template.docx"
NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}
ROLE_STYLES = (
    "Normal",
    "Figure",
    "Figure Caption",
    "Table Caption",
    "Equation",
    "Table",
    "Reference",
    "Heading 1",
    "Heading 2",
    "Heading 3",
)
FORBIDDEN_STYLES = (
    "SCI Body",
    "SCI Abstract Body",
    "SCI Statement Body",
    "SCI Three-Line Table",
    "图片",
    "图片标题",
    "表格标题",
)


class Audit:
    def __init__(self) -> None:
        self.results: list[dict[str, object]] = []

    def check(self, name: str, condition: bool, detail: str = "") -> None:
        self.results.append({"name": name, "status": "PASS" if condition else "FAIL", "detail": detail})

    @property
    def failures(self) -> int:
        return sum(item["status"] == "FAIL" for item in self.results)


def close(value, expected, tolerance=0.05) -> bool:
    return value is not None and abs(value - expected) <= tolerance


def point_value(length) -> float | None:
    return None if length is None else length.pt


def parse_part(archive: zipfile.ZipFile, name: str):
    return etree.fromstring(archive.read(name))


def style_is_exposed(style) -> bool:
    element = style._element
    return element.find(qn("w:qFormat")) is not None and element.find(qn("w:semiHidden")) is None


def style_by_id(document: Document, style_id: str):
    for style in document.styles:
        if style.style_id == style_id:
            return style
    raise KeyError(f"no style with id {style_id!r}")


def font_name_is_tnr(style) -> bool:
    font = style.font
    rpr = style._element.find(qn("w:rPr"))
    if rpr is None:
        return font.name == "Times New Roman"
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        return font.name == "Times New Roman"
    explicit_tnr = all(rfonts.get(qn(f"w:{attr}")) == "Times New Roman" for attr in ("ascii", "hAnsi"))
    no_theme_override = all(
        rfonts.get(qn(f"w:{attr}")) is None
        for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme")
    )
    return explicit_tnr and no_theme_override


def run_font_is_tnr(run) -> bool:
    rpr = run._r.find(qn("w:rPr"))
    if rpr is None:
        return False
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        return False
    return (
        all(rfonts.get(qn(f"w:{attr}")) == "Times New Roman" for attr in ("ascii", "hAnsi"))
        and all(
            rfonts.get(qn(f"w:{attr}")) is None
            for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme")
        )
    )


def zero_indents(fmt) -> bool:
    return all(
        close(point_value(value), 0)
        for value in (fmt.left_indent, fmt.right_indent, fmt.first_line_indent)
    )


def normal_indent_is_085_cm(fmt) -> bool:
    return (
        close(fmt.first_line_indent.cm if fmt.first_line_indent is not None else None, 0.85, 0.01)
        and close(point_value(fmt.left_indent), 0)
        and close(point_value(fmt.right_indent), 0)
    )


def audit_styles(document: Document, audit: Audit) -> None:
    names = {style.name for style in document.styles}
    audit.check("required manuscript-role styles exist", all(name in names for name in ROLE_STYLES), ", ".join(ROLE_STYLES))
    audit.check("legacy manuscript-role styles absent", not any(name in names for name in FORBIDDEN_STYLES), ", ".join(FORBIDDEN_STYLES))
    audit.check(
        "all manuscript-role styles exposed",
        all(style_is_exposed(document.styles[name]) for name in ROLE_STYLES),
        "ten English manuscript roles visible in Quick Styles",
    )

    normal = document.styles["Normal"]
    fmt = normal.paragraph_format
    audit.check("Normal font", font_name_is_tnr(normal) and close(point_value(normal.font.size), 12), "Times New Roman 12 pt")
    audit.check("Normal alignment and spacing", fmt.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY and fmt.line_spacing == 1.5, "justified, 1.5 spacing")
    audit.check("Normal paragraph spacing", close(point_value(fmt.space_before), 6) and close(point_value(fmt.space_after), 6), "6 pt before and after")
    audit.check("Normal body first-line indent", normal_indent_is_085_cm(fmt), "0.85 cm first-line; zero left and right indent")

    figure = document.styles["Figure"]
    audit.check(
        "Figure style",
        font_name_is_tnr(figure)
        and close(point_value(figure.font.size), 12)
        and figure.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
        and figure.paragraph_format.line_spacing == 1.0
        and close(point_value(figure.paragraph_format.space_before), 6)
        and close(point_value(figure.paragraph_format.space_after), 0)
        and zero_indents(figure.paragraph_format),
        "centered, single-spaced, zero indent",
    )

    figure_caption = document.styles["Figure Caption"]
    audit.check(
        "Figure Caption style",
        font_name_is_tnr(figure_caption)
        and close(point_value(figure_caption.font.size), 12)
        and figure_caption.font.italic is True
        and figure_caption.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
        and figure_caption.paragraph_format.line_spacing == 1.5
        and close(point_value(figure_caption.paragraph_format.space_before), 0)
        and close(point_value(figure_caption.paragraph_format.space_after), 6)
        and zero_indents(figure_caption.paragraph_format),
        "12 pt italic, justified, below figure",
    )

    table_caption = document.styles["Table Caption"]
    audit.check(
        "Table Caption style",
        font_name_is_tnr(table_caption)
        and close(point_value(table_caption.font.size), 12)
        and table_caption.font.italic is True
        and table_caption.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.LEFT
        and table_caption.paragraph_format.line_spacing == 1.5
        and close(point_value(table_caption.paragraph_format.space_before), 6)
        and close(point_value(table_caption.paragraph_format.space_after), 0)
        and zero_indents(table_caption.paragraph_format),
        "12 pt italic, left aligned, above table",
    )

    for level in range(1, 4):
        heading = document.styles[f"Heading {level}"]
        audit.check(
            f"Heading {level} style",
            font_name_is_tnr(heading)
            and close(point_value(heading.font.size), 14)
            and heading.font.bold is True
            and heading.paragraph_format.line_spacing == 1.5
            and zero_indents(heading.paragraph_format),
            "Times New Roman 14 pt bold",
        )

    equation = document.styles["Equation"]
    audit.check(
        "Equation style",
        font_name_is_tnr(equation)
        and close(point_value(equation.font.size), 12)
        and zero_indents(equation.paragraph_format),
        "Times New Roman 12 pt, zero indent",
    )

    table = document.styles["Table"]
    audit.check(
        "Table is a table style",
        table.type == WD_STYLE_TYPE.TABLE and font_name_is_tnr(table) and close(point_value(table.font.size), 10),
        "Times New Roman 10 pt table style",
    )

    line_number = style_by_id(document, "LineNumber")
    audit.check(
        "Line Number system style",
        line_number.type == WD_STYLE_TYPE.CHARACTER
        and font_name_is_tnr(line_number)
        and close(point_value(line_number.font.size), 11)
        and not style_is_exposed(line_number),
        "hidden system character style, Times New Roman 11 pt",
    )

    reference = document.styles["Reference"]
    audit.check(
        "Reference style",
        font_name_is_tnr(reference)
        and close(point_value(reference.font.size), 11)
        and reference.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.LEFT
        and reference.paragraph_format.line_spacing == 1.0
        and close(point_value(reference.paragraph_format.space_before), 0)
        and close(point_value(reference.paragraph_format.space_after), 3)
        and close(point_value(reference.paragraph_format.left_indent), 36)
        and close(point_value(reference.paragraph_format.first_line_indent), -36),
        "11 pt with 1.27 cm hanging indent",
    )


def audit_page(document: Document, document_xml, footer_xmls, header_xmls, settings_xml, audit: Audit) -> None:
    sections = list(document.sections)
    audit.check("single-column section exists", bool(sections), f"sections={len(sections)}")
    geometry_ok = all(
        close(section.page_width.cm, 21.0, 0.02)
        and close(section.page_height.cm, 29.7, 0.02)
        and all(
            close(value.cm, 2.54, 0.02)
            for value in (section.top_margin, section.bottom_margin, section.left_margin, section.right_margin)
        )
        and close(section.header_distance.cm, 1.27, 0.02)
        and close(section.footer_distance.cm, 1.27, 0.02)
        for section in sections
    )
    audit.check("A4 page geometry", geometry_ok, "21.0 x 29.7 cm; 2.54 cm margins; 1.27 cm header/footer")

    header_text = "".join("".join(xml.xpath("//w:t/text()", namespaces=NS)) for xml in header_xmls)
    audit.check("blank header", not header_text.strip(), f"header text={header_text!r}")

    sect_nodes = document_xml.xpath("//w:sectPr", namespaces=NS)
    line_number_nodes = document_xml.xpath("//w:sectPr/w:lnNumType", namespaces=NS)
    valid_line_numbers = len(line_number_nodes) == len(sect_nodes) and all(
        node.get(qn("w:countBy")) == "1"
        and node.get(qn("w:start")) == "0"
        and node.get(qn("w:restart")) == "continuous"
        for node in line_number_nodes
    )
    audit.check("continuous line numbering", valid_line_numbers, f"sections={len(sect_nodes)}, visible start=1")

    field_instructions = " ".join(
        text for xml in footer_xmls for text in xml.xpath("//w:instrText/text()", namespaces=NS)
    )
    audit.check("dynamic PAGE field", "PAGE" in field_instructions, field_instructions.strip())
    footer_centered = bool(footer_xmls) and all(
        xml.xpath("boolean(//w:p[w:fldChar or .//w:instrText]/w:pPr/w:jc[@w:val='center'])", namespaces=NS)
        for xml in footer_xmls
    )
    audit.check("footer PAGE field centered", footer_centered, "w:jc=center")

    update_fields = settings_xml.xpath("//w:updateFields", namespaces=NS)
    audit.check(
        "update fields on open",
        bool(update_fields) and update_fields[0].get(qn("w:val")) in {"true", "1"},
        "w:updateFields=true",
    )

    page_breaks = document_xml.xpath("//w:br[@w:type='page']", namespaces=NS)
    audit.check("front-matter page breaks", len(page_breaks) >= 2, f"page breaks={len(page_breaks)}")


def paragraph_style_id(paragraph_element) -> str | None:
    values = paragraph_element.xpath("./w:pPr/w:pStyle/@w:val", namespaces=NS)
    return values[0] if values else None


def audit_roles(document: Document, document_xml, audit: Audit) -> None:
    paragraph_roles = ("Normal", "Figure", "Figure Caption", "Table Caption", "Equation", "Reference", "Heading 1", "Heading 2", "Heading 3")
    used = document_xml.xpath("//w:p/w:pPr/w:pStyle/@w:val", namespaces=NS)
    for name in paragraph_roles:
        style_id = document.styles[name].style_id
        audit.check(f"{name} explicitly assigned", style_id in used, style_id)

    table_style_id = document.styles["Table"].style_id
    assigned_table_styles = document_xml.xpath("//w:tbl/w:tblPr/w:tblStyle/@w:val", namespaces=NS)
    audit.check("Table style explicitly assigned", table_style_id in assigned_table_styles, table_style_id)

    text_to_style = {paragraph.text.strip(): paragraph.style.name for paragraph in document.paragraphs}
    title_paragraphs = [p for p in document.paragraphs if p.text.startswith("Manuscript title:")]
    title_font_ok = bool(title_paragraphs) and all(
        run_font_is_tnr(run) for run in title_paragraphs[0].runs if run.text
    )
    audit.check("main title uses explicit Times New Roman", title_font_ok, "no theme-font override")
    audit.check("Abstract heading uses Heading 1", text_to_style.get("Abstract") == "Heading 1", str(text_to_style.get("Abstract")))
    abstract_paragraphs = [p for p in document.paragraphs if p.text.startswith("Replace this text with a self-contained abstract")]
    abstract_ok = (
        bool(abstract_paragraphs)
        and abstract_paragraphs[0].style.name == "Normal"
        and zero_indents(abstract_paragraphs[0].paragraph_format)
    )
    audit.check("abstract prose uses zero-indent Normal", abstract_ok, "Normal with direct zero-indent override")
    keyword_paragraphs = [p for p in document.paragraphs if p.text.startswith("Keywords:")]
    keyword_ok = bool(keyword_paragraphs) and keyword_paragraphs[0].style.name == "Normal" and keyword_paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    audit.check("Keywords uses flush-left Normal", keyword_ok, "Normal with direct left alignment")
    author_paragraphs = [p for p in document.paragraphs if p.text.startswith("First Author")]
    audit.check("author line centered", bool(author_paragraphs) and author_paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER, "centered")
    affiliation_paragraphs = [p for p in document.paragraphs if "Department, Institution" in p.text]
    audit.check(
        "affiliations flush left without indent",
        bool(affiliation_paragraphs)
        and all(p.alignment == WD_ALIGN_PARAGRAPH.LEFT and zero_indents(p.paragraph_format) for p in affiliation_paragraphs),
        f"count={len(affiliation_paragraphs)}",
    )
    declaration_headings = {"Acknowledgments", "Funding", "Competing interests", "Data availability", "Ethics statement", "CRediT authorship contribution statement"}
    declaration_bodies = []
    paragraphs = list(document.paragraphs)
    for index, paragraph in enumerate(paragraphs[:-1]):
        if paragraph.text in declaration_headings:
            declaration_bodies.append(paragraphs[index + 1])
    audit.check(
        "declaration prose uses Normal",
        len(declaration_bodies) == len(declaration_headings) and all(p.style.name == "Normal" for p in declaration_bodies),
        f"count={len(declaration_bodies)}",
    )

    body_children = list(document_xml.xpath("/w:document/w:body", namespaces=NS)[0])
    figure_style_id = document.styles["Figure"].style_id
    figure_caption_id = document.styles["Figure Caption"].style_id
    table_caption_id = document.styles["Table Caption"].style_id
    figure_pair = False
    table_pair = False
    inline_count = 0
    anchor_count = 0
    for index, child in enumerate(body_children):
        if child.tag == qn("w:p") and child.xpath(".//w:drawing", namespaces=NS):
            inline_count += len(child.xpath(".//wp:inline", namespaces=NS))
            anchor_count += len(child.xpath(".//wp:anchor", namespaces=NS))
            figure_pair = (
                paragraph_style_id(child) == figure_style_id
                and index + 1 < len(body_children)
                and paragraph_style_id(body_children[index + 1]) == figure_caption_id
            )
        if child.tag == qn("w:tbl") and index > 0:
            table_pair = paragraph_style_id(body_children[index - 1]) == table_caption_id
    audit.check("figure is inline", inline_count >= 1 and anchor_count == 0, f"inline={inline_count}, anchors={anchor_count}")
    audit.check("figure caption placement", figure_pair, "Figure Caption immediately below Figure")
    audit.check("table caption placement", table_pair, "Table Caption immediately above Table")


def audit_citations_and_references(document: Document, document_xml, audit: Audit) -> None:
    superscripts = document_xml.xpath("//w:r[w:rPr/w:vertAlign[@w:val='superscript']]/w:t/text()", namespaces=NS)
    audit.check("superscript numeric citation example", "1" in superscripts, f"superscripts={superscripts}")
    references = [p.text for p in document.paragraphs if p.style.name == "Reference"]
    order_ok = len(references) >= 2 and references[0].startswith("Zhao") and references[1].startswith("Adams")
    audit.check("reference source order preserved", order_ok, "Zhao remains before Adams; no alphabetization")


def audit_equation(document: Document, document_xml, audit: Audit) -> None:
    equation_style_id = document.styles["Equation"].style_id
    equation_paragraphs = document_xml.xpath(
        f"//w:p[w:pPr/w:pStyle[@w:val='{equation_style_id}']]", namespaces=NS
    )
    native_math = document_xml.xpath("//m:oMath", namespaces=NS)
    fractions = document_xml.xpath("//m:oMath//m:f", namespaces=NS)
    tab_nodes = sum(len(paragraph.xpath("./w:r/w:tab", namespaces=NS)) for paragraph in equation_paragraphs)
    style_tabs = document.styles["Equation"]._element.findall(".//" + qn("w:tab"))
    audit.check("native OMML equation", bool(native_math) and bool(fractions), f"oMath={len(native_math)}, fractions={len(fractions)}")
    audit.check("numbered equation tabs", tab_nodes >= 2 and len(style_tabs) >= 2, f"TAB nodes={tab_nodes}, style tab stops={len(style_tabs)}")


def audit_table(document_xml, audit: Audit) -> None:
    tables = document_xml.xpath("//w:tbl", namespaces=NS)
    audit.check("three-line table exists", bool(tables), f"tables={len(tables)}")
    if not tables:
        return
    table = tables[0]
    top = table.xpath("./w:tblPr/w:tblBorders/w:top", namespaces=NS)
    bottom = table.xpath("./w:tblPr/w:tblBorders/w:bottom", namespaces=NS)
    sides = table.xpath("./w:tblPr/w:tblBorders/w:left | ./w:tblPr/w:tblBorders/w:right | ./w:tblPr/w:tblBorders/w:insideH | ./w:tblPr/w:tblBorders/w:insideV", namespaces=NS)
    border_ok = (
        bool(top)
        and bool(bottom)
        and top[0].get(qn("w:sz")) == "8"
        and bottom[0].get(qn("w:sz")) == "8"
        and all(node.get(qn("w:val")) == "nil" for node in sides)
    )
    audit.check("three-line borders", border_ok, "1.0 pt top/bottom; no side or internal grid")

    header_cells = table.xpath("./w:tr[1]/w:tc", namespaces=NS)
    header_rule_ok = all(cell.xpath("./w:tcPr/w:tcBorders/w:bottom[@w:sz='4']", namespaces=NS) for cell in header_cells)
    repeat_header = bool(table.xpath("./w:tr[1]/w:trPr/w:tblHeader", namespaces=NS))
    audit.check("table header rule", header_rule_ok, "0.5 pt header separator")
    audit.check("table header repeats", repeat_header, "w:tblHeader")

    shading = table.xpath(".//w:shd", namespaces=NS)
    audit.check("table has no shading", not shading, f"shading nodes={len(shading)}")

    forbidden = ("firstLineChars", "hangingChars", "leftChars", "rightChars")
    bad_indents = [
        ind
        for ind in table.xpath(".//w:pPr/w:ind", namespaces=NS)
        if any(ind.get(qn(f"w:{name}")) is not None for name in forbidden)
    ]
    numeric_indent_ok = all(
        ind.get(qn("w:left")) == "0"
        and ind.get(qn("w:right")) == "0"
        and ind.get(qn("w:firstLine")) == "0"
        and ind.get(qn("w:hanging")) == "0"
        for ind in table.xpath(".//w:pPr/w:ind", namespaces=NS)
    )
    audit.check("table character indents removed", not bad_indents, f"bad indents={len(bad_indents)}")
    audit.check("table numeric indents zero", numeric_indent_ok, "left/right/firstLine/hanging=0")

    grid_widths = [int(value) for value in table.xpath("./w:tblGrid/w:gridCol/@w:w", namespaces=NS)]
    table_widths = [int(value) for value in table.xpath("./w:tr[1]/w:tc/w:tcPr/w:tcW/@w:w", namespaces=NS)]
    tbl_width = table.xpath("./w:tblPr/w:tblW/@w:w", namespaces=NS)
    geometry_ok = grid_widths == table_widths and bool(tbl_width) and sum(grid_widths) == int(tbl_width[0])
    audit.check("fixed table geometry", geometry_ok, f"grid={grid_widths}, total={tbl_width[0] if tbl_width else 'missing'}")

    cant_split = table.xpath("./w:tr/w:trPr/w:cantSplit", namespaces=NS)
    at_least = table.xpath("./w:tr/w:trPr/w:trHeight[@w:hRule='atLeast']", namespaces=NS)
    row_count = len(table.xpath("./w:tr", namespaces=NS))
    audit.check("table rows do not split", len(cant_split) == row_count, f"rows={row_count}, cantSplit={len(cant_split)}")
    audit.check("table row height is minimum", len(at_least) == row_count, f"atLeast={len(at_least)}")


def run_audit(path: Path) -> Audit:
    audit = Audit()
    audit.check("template exists", path.is_file(), str(path))
    if not path.is_file():
        return audit

    try:
        document = Document(path)
        audit.check("DOCX package opens", True, f"paragraphs={len(document.paragraphs)}, tables={len(document.tables)}")
    except Exception as exc:  # pragma: no cover
        audit.check("DOCX package opens", False, str(exc))
        return audit

    with zipfile.ZipFile(path) as archive:
        document_xml = parse_part(archive, "word/document.xml")
        settings_xml = parse_part(archive, "word/settings.xml")
        footer_names = [name for name in archive.namelist() if name.startswith("word/footer") and name.endswith(".xml")]
        header_names = [name for name in archive.namelist() if name.startswith("word/header") and name.endswith(".xml")]
        footer_xmls = [parse_part(archive, name) for name in footer_names]
        header_xmls = [parse_part(archive, name) for name in header_names]

    audit_styles(document, audit)
    audit_page(document, document_xml, footer_xmls, header_xmls, settings_xml, audit)
    audit_roles(document, document_xml, audit)
    audit_citations_and_references(document, document_xml, audit)
    audit_equation(document, document_xml, audit)
    audit_table(document_xml, audit)
    return audit


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("template", nargs="?", type=Path, default=DEFAULT_TEMPLATE)
    parser.add_argument("--json", action="store_true", dest="as_json")
    args = parser.parse_args()
    audit = run_audit(args.template.resolve())
    if args.as_json:
        print(json.dumps({"failures": audit.failures, "results": audit.results}, indent=2, ensure_ascii=False))
    else:
        for item in audit.results:
            detail = f" - {item['detail']}" if item["detail"] else ""
            print(f"{item['status']}: {item['name']}{detail}")
        print(f"SUMMARY: {len(audit.results) - audit.failures} PASS, {audit.failures} FAIL")
    sys.exit(1 if audit.failures else 0)


if __name__ == "__main__":
    main()
