#!/usr/bin/env python3
"""Build the reviewed SCI manuscript DOCX template."""

from __future__ import annotations

import argparse
from datetime import datetime, timezone
from pathlib import Path
from tempfile import TemporaryDirectory

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT = ROOT / "assets" / "SCI_manuscript_template.docx"
BLACK = RGBColor(0, 0, 0)
BLUE = RGBColor(5, 99, 193)
WRITABLE_WIDTH_CM = 15.92


def remove_child(parent, tag: str) -> None:
    for child in list(parent):
        if child.tag == qn(tag):
            parent.remove(child)


def set_font_name(font, name: str = "Times New Roman") -> None:
    font.name = name
    rpr = font._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        key = qn(f"w:{attr}")
        if key in rfonts.attrib:
            del rfonts.attrib[key]


def set_run_font(run, size: float, *, bold=None, italic=None, color=BLACK) -> None:
    set_font_name(run.font)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def expose_style(style) -> None:
    element = style._element
    for tag in ("w:semiHidden", "w:unhideWhenUsed"):
        remove_child(element, tag)
    if element.find(qn("w:qFormat")) is None:
        element.append(OxmlElement("w:qFormat"))


def hide_style(style) -> None:
    element = style._element
    remove_child(element, "w:qFormat")
    if element.find(qn("w:semiHidden")) is None:
        element.append(OxmlElement("w:semiHidden"))


def normalize_line_number_style(style) -> None:
    """Encode Line Number as Word's built-in style so Word preserves its rPr."""
    element = style._element
    element.attrib.pop(qn("w:customStyle"), None)
    name = element.find(qn("w:name"))
    if name is not None:
        name.set(qn("w:val"), "line number")
    based_on = element.find(qn("w:basedOn"))
    if based_on is None:
        based_on = OxmlElement("w:basedOn")
        element.insert(1, based_on)
    based_on.set(qn("w:val"), "DefaultParagraphFont")
    priority = element.find(qn("w:uiPriority"))
    if priority is None:
        priority = OxmlElement("w:uiPriority")
        semi_hidden = element.find(qn("w:semiHidden"))
        element.insert(element.index(semi_hidden) if semi_hidden is not None else 2, priority)
    priority.set(qn("w:val"), "99")
    if element.find(qn("w:unhideWhenUsed")) is None:
        unhide = OxmlElement("w:unhideWhenUsed")
        semi_hidden = element.find(qn("w:semiHidden"))
        if semi_hidden is not None:
            element.insert(element.index(semi_hidden) + 1, unhide)
        else:
            element.append(unhide)


def get_or_add_style(document: Document, name: str, style_type, *, expose: bool = True):
    try:
        style = document.styles[name]
    except KeyError:
        style = document.styles.add_style(name, style_type)
    if expose:
        expose_style(style)
    else:
        hide_style(style)
    return style


def set_zero_indent(target) -> None:
    fmt = target.paragraph_format if hasattr(target, "paragraph_format") else target
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)
    fmt.first_line_indent = Pt(0)
    element = target._p if hasattr(target, "_p") else target._element
    ppr = element.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    for attr in ("left", "right", "firstLine", "hanging"):
        ind.set(qn(f"w:{attr}"), "0")
    for attr in ("leftChars", "rightChars", "firstLineChars", "hangingChars"):
        key = qn(f"w:{attr}")
        if key in ind.attrib:
            del ind.attrib[key]


def configure_paragraph_style(
    document: Document,
    name: str,
    *,
    size: float,
    alignment,
    line_spacing: float,
    before: float,
    after: float,
    bold: bool = False,
    italic: bool = False,
    outline_level: int | None = None,
    first_line_cm: float = 0,
):
    style = get_or_add_style(document, name, WD_STYLE_TYPE.PARAGRAPH)
    set_font_name(style.font)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = BLACK
    fmt = style.paragraph_format
    fmt.alignment = alignment
    fmt.line_spacing = line_spacing
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    set_zero_indent(style)
    if first_line_cm:
        fmt.first_line_indent = Cm(first_line_cm)
    if outline_level is not None:
        ppr = style._element.get_or_add_pPr()
        outline = ppr.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            ppr.append(outline)
        outline.set(qn("w:val"), str(outline_level))
    return style


def configure_styles(document: Document) -> None:
    configure_paragraph_style(
        document,
        "Normal",
        size=12,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        line_spacing=1.5,
        before=6,
        after=6,
        first_line_cm=0.85,
    )
    configure_paragraph_style(
        document,
        "Figure",
        size=12,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.0,
        before=6,
        after=0,
    )
    configure_paragraph_style(
        document,
        "Figure Caption",
        size=12,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        line_spacing=1.5,
        before=0,
        after=6,
        italic=True,
    )
    configure_paragraph_style(
        document,
        "Table Caption",
        size=12,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.5,
        before=6,
        after=0,
        italic=True,
    )
    equation = configure_paragraph_style(
        document,
        "Equation",
        size=12,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.0,
        before=6,
        after=6,
    )
    equation.paragraph_format.tab_stops.clear_all()
    equation.paragraph_format.tab_stops.add_tab_stop(Cm(WRITABLE_WIDTH_CM / 2), WD_TAB_ALIGNMENT.CENTER)
    equation.paragraph_format.tab_stops.add_tab_stop(Cm(WRITABLE_WIDTH_CM), WD_TAB_ALIGNMENT.RIGHT)

    reference = configure_paragraph_style(
        document,
        "Reference",
        size=11,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.0,
        before=0,
        after=3,
    )
    reference.paragraph_format.left_indent = Cm(1.27)
    reference.paragraph_format.first_line_indent = Cm(-1.27)

    for level in range(1, 4):
        configure_paragraph_style(
            document,
            f"Heading {level}",
            size=14,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            line_spacing=1.5,
            before=6,
            after=0,
            bold=True,
            outline_level=level - 1,
        )

    table_style = get_or_add_style(document, "Table", WD_STYLE_TYPE.TABLE)
    set_font_name(table_style.font)
    table_style.font.size = Pt(10)
    table_style.font.color.rgb = BLACK

    line_number = get_or_add_style(
        document, "Line Number", WD_STYLE_TYPE.CHARACTER, expose=False
    )
    set_font_name(line_number.font)
    line_number.font.size = Pt(11)
    line_number.font.bold = False
    line_number.font.italic = False
    line_number.font.color.rgb = BLACK
    normalize_line_number_style(line_number)


def append_page_field(paragraph) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    cached = OxmlElement("w:t")
    cached.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    set_run_font(run, 12)
    run._r.extend((begin, instruction, separate, cached, end))


def configure_section(section) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.27)
    section.footer_distance = Cm(1.27)

    header = section.header
    for paragraph in header.paragraphs:
        paragraph.clear()

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_zero_indent(paragraph)
    append_page_field(paragraph)

    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:lnNumType"))
    if existing is not None:
        sect_pr.remove(existing)
    line_numbers = OxmlElement("w:lnNumType")
    line_numbers.set(qn("w:countBy"), "1")
    line_numbers.set(qn("w:start"), "0")
    line_numbers.set(qn("w:restart"), "continuous")
    line_numbers.set(qn("w:distance"), "360")
    sect_pr.append(line_numbers)


def set_update_fields(document: Document) -> None:
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def force_explicit_normal_styles(document: Document) -> None:
    """Keep an explicit w:pStyle=Normal marker on every main-story Normal paragraph."""
    normal_id = document.styles["Normal"].style_id
    for paragraph in document.paragraphs:
        if paragraph.style.name != "Normal":
            continue
        ppr = paragraph._p.get_or_add_pPr()
        pstyle = ppr.find(qn("w:pStyle"))
        if pstyle is None:
            pstyle = OxmlElement("w:pStyle")
            ppr.insert(0, pstyle)
        pstyle.set(qn("w:val"), normal_id)


def format_front_paragraph(paragraph, *, size: float, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    paragraph.alignment = alignment
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(0)
    set_zero_indent(paragraph)
    for run in paragraph.runs:
        set_run_font(run, size, bold=bold)


def add_title_block(document: Document) -> None:
    title = document.add_paragraph(style="Normal")
    title.add_run("Manuscript title: replace with a concise, informative title")
    format_front_paragraph(title, size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    authors = document.add_paragraph(style="Normal")
    authors.add_run("First Author").bold = False
    marker1 = authors.add_run("1")
    marker1.font.superscript = True
    authors.add_run(", Second Author")
    marker2 = authors.add_run("2")
    marker2.font.superscript = True
    format_front_paragraph(authors, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    affiliation1 = document.add_paragraph(style="Normal")
    sup1 = affiliation1.add_run("1")
    sup1.font.superscript = True
    affiliation1.add_run(" Department, Institution, City, Postal code, Country")
    format_front_paragraph(affiliation1, size=12)

    affiliation2 = document.add_paragraph(style="Normal")
    sup2 = affiliation2.add_run("2")
    sup2.font.superscript = True
    affiliation2.add_run(" Department, Institution, City, Postal code, Country")
    format_front_paragraph(affiliation2, size=12)

    corresponding = document.add_paragraph(style="Normal")
    corresponding.add_run("* Corresponding author: ")
    email = corresponding.add_run("author@example.com")
    set_run_font(email, 12, color=BLUE)
    email.underline = True
    format_front_paragraph(corresponding, size=12)
    corresponding.add_run().add_break(WD_BREAK.PAGE)


def add_abstract(document: Document) -> None:
    document.add_paragraph("Abstract", style="Heading 1")
    abstract = document.add_paragraph(
        "Replace this text with a self-contained abstract. This paragraph uses the Normal style: Times New Roman 12 pt, justified, 1.5 line spacing, 6 pt before and after, and no first-line indent.",
        style="Normal",
    )
    set_zero_indent(abstract)
    keywords = document.add_paragraph(style="Normal")
    keywords.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_zero_indent(keywords)
    label = keywords.add_run("Keywords: ")
    set_run_font(label, 12, bold=True)
    set_run_font(keywords.add_run("landslide; physical model; monitoring; deformation; uncertainty"), 12)
    keywords.add_run().add_break(WD_BREAK.PAGE)


def create_placeholder_figure(path: Path) -> None:
    image = Image.new("RGB", (1500, 720), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((80, 70, 1420, 650), outline=(65, 65, 65), width=4)
    draw.line((150, 585, 1360, 585), fill=(65, 65, 65), width=3)
    draw.line((150, 585, 150, 130), fill=(65, 65, 65), width=3)
    points = [(150, 520), (350, 490), (550, 410), (750, 330), (950, 260), (1150, 210), (1360, 170)]
    draw.line(points, fill=(55, 103, 168), width=8)
    for x, y in points:
        draw.ellipse((x - 9, y - 9, x + 9, y + 9), fill=(55, 103, 168))
    draw.text((570, 90), "Replace with final figure", fill=(30, 30, 30))
    image.save(path, format="PNG", optimize=True)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single" if edge in {"top", "bottom"} else "nil")
        node.set(qn("w:color"), "000000")
        node.set(qn("w:sz"), "8" if edge in {"top", "bottom"} else "0")


def set_cell_bottom_border(cell, size: int = 4) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:color"), "000000")


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, twips in (("top", 60), ("bottom", 60), ("left", 80), ("right", 80)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(twips))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    width = tc_pr.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        tc_pr.append(width)
    width.set(qn("w:w"), str(twips))
    width.set(qn("w:type"), "dxa")


def configure_table(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    tbl_pr = table._tbl.tblPr
    table_width = tbl_pr.find(qn("w:tblW"))
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(sum(widths)))
    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "0")
    indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row_index, row in enumerate(table.rows):
        row.height = Cm(0.7)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        tr_pr = row._tr.get_or_add_trPr()
        no_split = OxmlElement("w:cantSplit")
        tr_pr.append(no_split)
        if row_index == 0:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
        for column_index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[column_index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_index == 0:
                set_cell_bottom_border(cell, 4)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                    if row_index > 0 and column_index == 1
                    else WD_ALIGN_PARAGRAPH.CENTER
                )
                set_zero_indent(paragraph)
                for run in paragraph.runs:
                    set_run_font(run, 10, bold=(row_index == 0))


def omath_run(text: str, *, italic: bool = True):
    run = OxmlElement("m:r")
    if italic:
        rpr = OxmlElement("m:rPr")
        style = OxmlElement("m:sty")
        style.set(qn("m:val"), "i")
        rpr.append(style)
        run.append(rpr)
    value = OxmlElement("m:t")
    value.text = text
    run.append(value)
    return run


def build_native_equation():
    equation = OxmlElement("m:oMath")
    equation.append(omath_run("v"))
    equation.append(omath_run("=", italic=False))
    fraction = OxmlElement("m:f")
    numerator = OxmlElement("m:num")
    numerator.append(omath_run("Δd"))
    denominator = OxmlElement("m:den")
    denominator.append(omath_run("Δt"))
    fraction.append(numerator)
    fraction.append(denominator)
    equation.append(fraction)
    return equation


def add_equation(document: Document) -> None:
    paragraph = document.add_paragraph(style="Equation")
    set_run_font(paragraph.add_run("\t"), 12)
    paragraph._p.append(build_native_equation())
    set_run_font(paragraph.add_run("\t"), 12)
    set_run_font(paragraph.add_run("(1)"), 12)


def add_manuscript_elements(document: Document, figure_path: Path) -> None:
    document.add_paragraph("1. Introduction", style="Heading 1")
    paragraph = document.add_paragraph(style="Normal")
    set_run_font(
        paragraph.add_run(
            "Replace with manuscript text. Preserve an author-year system such as Smith (2024), or an existing numeric citation such as "
        ),
        12,
    )
    citation = paragraph.add_run("1")
    set_run_font(citation, 12)
    citation.font.superscript = True
    set_run_font(paragraph.add_run("."), 12)

    document.add_paragraph("1.1 Study context", style="Heading 2")
    document.add_paragraph(
        "Use Normal for ordinary narrative paragraphs.",
        style="Normal",
    )
    document.add_paragraph("1.1.1 Optional third-level heading", style="Heading 3")
    document.add_paragraph(
        "All three heading levels use Times New Roman.",
        style="Normal",
    )

    document.add_paragraph("2. Figures", style="Heading 1")
    figure_paragraph = document.add_paragraph(style="Figure")
    figure_paragraph.add_run().add_picture(str(figure_path), width=Cm(WRITABLE_WIDTH_CM))
    document.add_paragraph(
        "Figure 1. Replace this placeholder with the final inline figure and retain this 12 pt italic, justified caption below the image.",
        style="Figure Caption",
    )

    document.add_paragraph("3. Tables", style="Heading 1")
    document.add_paragraph(
        "Table 1. Illustrative physical-model test conditions and observed responses.",
        style="Table Caption",
    )
    table = document.add_table(rows=4, cols=4, style="Table")
    values = [
        ["Model ID", "Failure mode", "Peak displacement (mm)", "Failure time (min)"],
        ["PM-01", "No failure", "8.4", ">180"],
        ["PM-02", "Progressive sliding", "18.7", "146"],
        ["PM-03", "Toe failure", "36.2", "93"],
    ]
    for row, row_values in zip(table.rows, values):
        for cell, value in zip(row.cells, row_values):
            cell.text = value
    configure_table(table, [1500, 3000, 2300, 2226])

    document.add_paragraph("4. Equations", style="Heading 1")
    document.add_paragraph(
        "Use native editable Word equations for mathematical expressions that require equation typesetting.",
        style="Normal",
    )
    add_equation(document)

    document.add_paragraph("5. References", style="Heading 1")
    document.add_paragraph(
        "Zhao, B. (2022). First source-position example; preserve this position and its citation mapping.",
        style="Reference",
    )
    document.add_paragraph(
        "Adams, C. (2024). Second source-position example; do not alphabetize, reorder, or renumber automatically.",
        style="Reference",
    )

    for heading, text in (
        ("Acknowledgments", "Replace with acknowledgments."),
        ("Funding", "Replace with funding information."),
        ("Competing interests", "Replace with a competing-interest statement."),
        ("Data availability", "Replace with a data-availability statement."),
        ("Ethics statement", "Replace with an ethics statement when applicable."),
    ):
        document.add_paragraph(heading, style="Heading 1")
        document.add_paragraph(text, style="Normal")

    document.add_paragraph("CRediT authorship contribution statement", style="Heading 1")
    credit = document.add_paragraph(style="Normal")
    name = credit.add_run("First Author: ")
    set_run_font(name, 12, bold=True)
    set_run_font(credit.add_run("Conceptualization, Methodology, Writing – original draft."), 12)


def build_template(output: Path) -> Path:
    output.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_styles(document)
    for section in document.sections:
        configure_section(section)
    set_update_fields(document)

    properties = document.core_properties
    properties.title = "SCI Manuscript Template"
    properties.subject = "Reviewed English-language SCI manuscript formatting template"
    properties.author = "SCI Manuscript Format skill"
    properties.keywords = "SCI manuscript, DOCX template, three-line table, OMML"
    properties.created = datetime(2026, 9, 1, tzinfo=timezone.utc)
    properties.modified = datetime(2026, 9, 1, tzinfo=timezone.utc)

    add_title_block(document)
    add_abstract(document)
    with TemporaryDirectory(prefix="sci-template-") as temporary:
        figure_path = Path(temporary) / "figure-placeholder.png"
        create_placeholder_figure(figure_path)
        add_manuscript_elements(document, figure_path)

    force_explicit_normal_styles(document)
    document.save(output)
    return output


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    path = build_template(args.output.resolve())
    print(path)


if __name__ == "__main__":
    main()
